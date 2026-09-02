#!/usr/bin/env python3
"""生成毕业论文大纲 v2 .docx —— 整合所有已完成和待完成工作"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

OUT = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                   "毕业论文大纲_v2_自我优势效应的实验设计空间优化.docx")
doc = Document()
style = doc.styles['Normal']
style.font.name = '宋体'; style.font.size = Pt(12)
style.paragraph_format.space_after = Pt(6); style.paragraph_format.line_spacing = 1.5

def H(level, text):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = '黑体'; r.font.color.rgb = RGBColor(0,0,0)
        r.font.size = {1:16,2:14,3:12}.get(level,12)
    return h

def P(text, bold=False):
    p = doc.add_paragraph()
    if bold: p.add_run(text).bold = True
    else: p.add_run(text)
    return p

def Bullets(items):
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

def Table(headers, rows):
    t = doc.add_table(rows=len(rows)+1, cols=len(headers), style='Light Grid Accent 1')
    for i,h in enumerate(headers):
        t.rows[0].cells[i].text = h
        for p in t.rows[0].cells[i].paragraphs:
            for r in p.runs: r.bold = True; r.font.size = Pt(9)
    for rr,row in enumerate(rows):
        for cc,val in enumerate(row):
            t.rows[rr+1].cells[cc].text = str(val)
            for p in t.rows[rr+1].cells[cc].paragraphs:
                for r in p.runs: r.font.size = Pt(9)

# ═══ 封面 ═══
for _ in range(3): doc.add_paragraph()
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('硕 士 学 位 论 文 大 纲（v2.0）'); r.font.size=Pt(22); r.font.name='黑体'; r.bold=True
doc.add_paragraph()
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('自我优势效应的实验设计空间优化：\n基于DDM的参数预测与多源验证'); r.font.size=Pt(18); r.font.name='黑体'; r.bold=True
for _ in range(3): doc.add_paragraph()
doc.add_paragraph().add_run('— 2026年 —').font.size=Pt(14)
doc.add_page_break()

# ═══ 第1章 ═══
H(1,'第1章  绪论')
H(2,'1.1  研究背景')
P('自我优势效应（SPE）是指个体对自我相关刺激加工更快更准的现象（Sui et al., 2012）。在Self-Matching Task（SMT）范式中，实验参数——练习次数P、刺激呈现时间T与反应窗口W——可以系统性操控。JH（前人）提出了基于Sigmoid函数和DDM的设计空间建模方法，但其模型(1)缺乏对预测不确定性的量化，(2)未处理omission试次的影响，(3)未在更大范围数据上进行外部验证。本研究在此基础上引入高斯过程（GP）作为残差捕捉工具，将omission纳入显式建模，并通过大规模SPE数据库进行多源验证。')

H(2,'1.2  理论框架：从实验参数到心理过程的映射')
P('本研究基于四条理论假设（详见《理论逻辑链评估与修正建议》）：', bold=True)
Bullets([
    '假设一 (P→v)：练习次数通过v_P(P)影响漂移率——练习越多，信息累积越快',
    '假设二 (T→v)：刺激呈现时间通过v_T(T)影响漂移率——T越短，信息质量越差，v越低。选择将T映射到v而非Ter是一个有条件且可检验的建模选择',
    '假设三 (M=T+W→a)：总可用时间通过a(M)影响决策边界——时间压力越大，决策标准越低。T在模型中扮演双重角色：既影响v（信息质量），也通过M影响a（时间压力）',
    '假设四 (Sigmoid形式)：上述三条映射均采用Sigmoid函数（S型曲线），反映心理物理学的阈值-增长-饱和三阶段模式（Luce, 1986）'
])

H(2,'1.3  研究问题')
P('核心研究问题：实验设计参数(P,T,W)如何通过改变潜在心理决策过程(DDM参数v,a,t,z)，进而影响心理现象(SPE)的表现形式和稳定性？')
P('具体研究问题：')
Bullets([
    'RQ1 (方法论)：Omission试次的不同处理方法如何偏倚DDM参数估计？有无一个可接受的遗漏率阈值？',
    'RQ2 (建模)：能否构建一个GP+Sigmoid混合生成模型，从(P,T,W)预测DDM参数(v,a,t,z)，并在行为层面重建真实被试的SPE模式？',
    'RQ3 (Omission扩展)：能否将omission概率显式纳入生成模型，使其同时预测RT、ACC和omission率？',
    'RQ4 (外部验证)：模型预测能否在大规模跨研究数据（SPE Database, 44篇/3603人）和CRF仿真中得到外部效度支持？'
])

H(2,'1.4  研究意义')
P('理论意义：(1)提出了Sigmoid(理论刚性)+GP(数据柔性)的混合建模框架；(2)通过Censor vs Drop系统对比为omission处理方法选择提供了实证依据；(3)首次将LAN+OPN方法应用于SMT范式。')
P('实践意义：(1)为SMT实验设计提供了量化的参数选择指导；(2)基于GP不确定性的候选实验点推荐降低了试错成本。')

H(2,'1.5  论文结构')
P('本文共七章。第1-2章为绪论与文献综述。第3-6章包含四项实证研究。第7章为总讨论。')
doc.add_page_break()

# ═══ 第2章 ═══
H(1,'第2章  文献综述')
H(2,'2.1  Self-Matching Task与自我优势效应')
P('综述Sui et al.(2012)的SMT范式、SPE的行为特征及影响因素。特别梳理实验设计参数(P,T,W)如何被已有文献操纵。')

H(2,'2.2  漂移扩散模型（DDM）')
H(3,'2.2.1  DDM基本理论与参数')
P('综述Ratcliff(1978)及Ratcliff & McKoon(2008)的理论框架。核心参数v,a,z,Ter及其心理学含义。')
H(3,'2.2.2  DDM参数的系统性综述')
P('基于Tran et al.(2021,158篇)和Matzke & Wagenmakers(2009,23篇)的参数经验分布与推荐先验。为本研究参数合理性评估提供参照。')
H(3,'2.2.3  DDM中的Omission问题')
P('综述Leng et al.(2025)的LAN+OPN框架——丢弃omission会严重偏倚参数估计。三种方案：丢弃(Drop)、截尾(Censor)、显式建模(OPN)。')

H(2,'2.3  高斯过程（GP）与实验设计优化')
P('综述Schulz et al.(2018)的GP回归教程、Myung et al.(2013)的最优实验设计框架及贝叶斯主动学习策略。')

H(2,'2.4  SPE Database与大规模数据整合')
P('综述SPE Database（蔡振辛等，44篇/70数据集/3603人）的三层标准化框架、五组件任务分解及三项大规模分析发现（基线条件对SPE的影响、Mismatch下的小SPE、刺激呈现时间与SPE的关系）。')

H(2,'2.5  CRF方法与DDM参数推断')
P('综述条件反应函数（CRF）在决策偏差研究中的应用。简述从CRF曲线反推DDM参数的理论可行性及参数间的补偿效应。')

H(2,'2.6  文献小结与研究空白')
P('现有文献缺乏：(1)从实验参数到DDM参数的系统性映射框架；(2)omission处理方法的定量影响评估来自真实数据；(3)多源外部验证的建模论文。本研究旨在填补这三项空白。')
doc.add_page_break()

# ═══ 第3章 ═══
H(1,'第3章  研究一：SPE设计空间基本效应与底层假设检验')
P('（整合 ANOVA/BF分析 + SPE Database设计空间探索 + CRF可视化）')

H(2,'3.1  研究问题')
P('RQ1a: 8组实验条件间是否存在SPE差异？RQ1b: SPE Database中设计空间参数如何影响SPE的表现形式？')

H(2,'3.2  方法')
H(3,'3.2.1  数据来源')
Bullets([
    '数据集A (内部): 88名被试×8组实验条件（课题组前期采集），包含260 trials/人的Matching试次',
    '数据集B (外部): SPE Database——44篇文献/70数据集/3603名被试/155万+试次',
    '数据集C (仿真): CRF仿真数据（Euler-Maruyama数值积分 + HDDM生成器，z-bias 4水平×30被试）'
])
H(3,'3.2.2  分析策略')
Bullets([
    '内部验证: 单因素ANOVA + G*Power敏感性分析 + 贝叶斯因子(BF₁₀, JZS先验)',
    '外部验证: 跨研究SPE分布对比 + RT窗口分析(滑动窗口100-400ms) + 设计空间变量(刺激呈现时间、试次数)与SPE的关系',
    'CRF分析: 条件反应函数曲线可视化 + DDM参数→CRF逆向推断可行性评估'
])

H(2,'3.3  结果')
P('（填入ANOVA结果表、BF表、SPE Database关键图表、CRF曲线图）')

H(2,'3.4  讨论')
P('(1)ANOVA支持条件间存在差异(p=.019)，但线性回归无法解释(R²=.055)——引出非线性建模的必要性。'
  '(2)SPE Database发现：SPE在RT 0.3-0.8s范围内最强，Mismatch下存在小但稳定的SPE，刺激呈现时间对SPE有调制效应——这些外部发现为本研究的GP+Sigmoid建模提供了独立的行为参照。'
  '(3)CRF仿真建立了z-bias→行为模式的生成链，为DDM参数可解释性提供了机制层面证据。')
doc.add_page_break()

# ═══ 第4章 ═══
H(1,'第4章  研究二：Omission处理方法对DDM参数估计的系统影响')
P('（对应 Omission_Sensitivity_Analysis.ipynb + OPN_Training 第一版代码）')

H(2,'4.1  研究问题')
P('RQ2a: Censor(截尾)方案和Drop(直接丢弃)方案下DDM参数估计是否存在系统性差异？RQ2b: LAN+OPN框架能否进一步改善参数恢复？')

H(2,'4.2  方法')
H(3,'4.2.1  Censor vs Drop敏感性分析（已完成）')
Bullets([
    '设计: 8组×2方案=16次独立HDDM层次模型拟合',
    'Censor方案: omission试次rt=T+W, response=0, p_outlier=0',
    'Drop方案: 直接删除omission试次, p_outlier=0.05',
    '评估: 各参数差异Δ、95%CI重叠判断、Cohen\'s d'
])
H(3,'4.2.2  OPN训练框架（第一版已完成，待端到端验证）')
Bullets([
    'DDM仿真器: 批量版本simulate_ddm_with_deadline()',
    '训练数据生成: (v,a,t,z,deadline)→omission概率标签',
    '网络架构: MLP Regressor (sklearn)',
    '目标: OPN输入(θ, deadline)，输出log-probability of omission'
])
H(3,'4.2.3  LAN+OPN联合似然路线图（方法论已制定）')
Bullets([
    '简化策略: 用HDDM解析似然替代LAN——只需训练OPN',
    '联合似然: log_l = Σ log hddm_lik(observed) + |O| × OPN(θ, deadline)',
    'MCMC: PyMC层级贝叶斯模型实现参数估计'
])

H(2,'4.3  结果')
P('（填入Censor vs Drop对比表、OPN训练收敛曲线、LAN+OPN参数恢复结果）')

H(2,'4.4  讨论')
P('(1)G1(遗漏率72%)的v_self在两种方案间符号反转(Δ=6.81)，证实Leng et al.(2025)偏倚机制。'
  '(2)遗漏率>35%: 偏倚不可接受(95%CI不重叠); 遗漏率<15%: Censor方案足够可靠。'
  '(3)OPN框架验证了Leng et al.的方法论在本实验范式中的适用性。'
  '(4)据此，后续GP+Sigmoid建模排除G1/G2，保留G3-G8。')
doc.add_page_break()

# ═══ 第5章 ═══
H(1,'第5章  研究三：GP+Sigmoid混合生成模型的构建与验证')
P('（对应 GP+Sigmoid Cleaned Pipeline）')

H(2,'5.1  研究问题')
P('RQ3: 能否构建一个结合Sigmoid理论先验与GP残差捕捉的混合生成模型，实现对任意(P,T,W)组合下DDM参数的预测？')

H(2,'5.2  方法')
H(3,'5.2.1  数据筛选')
P('基于研究二结论，排除G1/G2（遗漏率>50%）。G3-G4标记为谨慎使用，G5-G8作为核心建模样本（N=6组）。')
H(3,'5.2.2  Sigmoid理论参数化')
P('v(T,P,condition)=v_T(T)×v_P(P)×base_scale_v×(1+condition_modulation)。'
  'a(M)=Sigmoid(M)×base_scale_a×(1+boundary_modulation)。'
  'T双重角色的理论说明：T同时出现在v_T(T)和a(M)=a(T+W)中，反映T对信息质量和时间压力的双重效应。')
H(3,'5.2.3  Sigmoid参数校准')
P('差分进化算法(Differential Evolution)最小化Sigmoid预测与HDDM真实参数的RMSE。'
  '7个可优化参数：alaph1, alaph2, beta1, beta2, gamma, base_scale_v, base_scale_a。')
H(3,'5.2.4  GP+Sigmoid混合模型')
P('GPSigmoidHybridModel: 5个独立GP(v_self, v_stranger, a, t, z)，核函数ConstantKernel×RBF+WhiteKernel。'
  'GP学习目标：真实HDDM参数 - Sigmoid预测值。最终预测 = Sigmoid + GP残差。')
H(3,'5.2.5  验证策略')
Bullets([
    'In-sample训练拟合: RMSE和r',
    'LOCV交叉验证: Leave-One-Condition-Out',
    '行为层面验证: 用GP+Sigmoid预测参数进行DDM试次级模拟→对比模拟vs真实行为(RT, ACC, Omission, SPE)'
])

H(2,'5.3  结果')
P('（填入Sigmoid校准参数表、训练拟合指标、LOCV指标、行为验证指标）')

H(2,'5.4  讨论')
P('主要发现：(1)alaph1从默认1.5降至0.199——self的相对优势仅+20%而非假设的+150%; '
  '(2)beta1符号反转——高M条件下边界反而降低; '
  '(3)LOCV: v r=-0.09——6个点不足以支持GP泛化; '
  '(4)行为验证: RT r=0.98, ACC r=0.97——Sigmoid理论先验对行为重建起保障作用; '
  '(5)GP不确定性最高的区域(T=500ms, W=300ms附近)被推荐为下一轮实验条件的候选区域。')
doc.add_page_break()

# ═══ 第6章 ═══
H(1,'第6章  研究四：CRF与Stim-Coding仿真——DDM参数可解释性的机制探索')
P('（整合 CRF分析 + Stim-Coding仿真 + DDM参数→CRF逆向推断）')

H(2,'6.1  研究问题')
P('RQ4a: DDM参数(v,a,t,z)如何系统地影响条件反应函数(CRF)曲线的形态？'
  'RQ4b: 能否从实测CRF曲线反推底层DDM参数？（逆向推断可行性）')

H(2,'6.2  方法')
H(3,'6.2.1  CRF数据来源')
Bullets([
    '真实CRF数据: 88名被试的Matching试次RT和ACC（来自内部数据集A）',
    'CRF计算: 分位数条件反应函数——按RT分位数计算各条件下的正确率',
    '可视化: V1-V4共四版迭代（Matplotlib→Plotly交互→Shiny Web→综合R分析）'
])
H(3,'6.2.2  Stim-Coding仿真')
Bullets([
    '范式: Stim-Coding框架——通过操控DDM起始点z模拟自我-他人偏差',
    '4个z水平: neutral(0.50), small(0.55), medium(0.60), large(0.65)',
    '双引擎: Euler-Maruyama数值积分(Wiener模式) + HDDM内置生成器(HDDM模式)',
    '30被试×4条件×150试次=18,000仿真试次'
])
H(3,'6.2.3  DDM参数→CRF逆向推断')
Bullets([
    '正向: 设定(v,a,t,z)→仿真→CRF曲线',
    '逆向(目标): 实测CRF→优化→估计(v̂,â,t̂,ẑ)',
    '方法: 最小化Loss(CRF_simulated(v,a,t,z), CRF_observed)',
    '参数可辨识性: 4个参数对CRF有不同且可分离的影响(v→高度, a→陡峭度, z→偏移, t0→平移)'
])

H(2,'6.3  结果')
P('（填入CRF曲线图、Stim-Coding仿真CRF与实测CRF对比、逆向推断参数恢复精度）')

H(2,'6.4  讨论')
P('(1)Stim-Coding仿真生成的CRF模式与真实被试的CRF模式定性一致——支持DDM参数作为SPE心理机制的合理模型。'
  '(2)逆向推断面临v-z补偿效应和a-v部分补偿的挑战——参数可辨识性在统计上有限但在定性层面有信息。'
  '(3)本研究的CRF分析为GP+Sigmoid模型的参数预测提供了机制层面的"跨范式验证"。')
doc.add_page_break()

# ═══ 第7章 ═══
H(1,'第7章  总讨论')
H(2,'7.1  研究发现总结')
Bullets([
    '(1)Omission处理：遗漏率>35%时参数偏倚不可接受，<15%时Censor方案可靠——为DDM分析方法论提供了实证阈值（研究二）',
    '(2)GP+Sigmoid混合模型：行为层面重建优异(RT r=0.98)，参数层面泛化受限于样本量，但Sigmoid理论先验提供了必要约束（研究三）',
    '(3)多源外部验证：SPE Database发现为模型提供了跨范式、跨人群的外部效度支持（研究一、研究四）',
    '(4)Sigmoid校准揭示了非平凡的理论参数变化(alaph1≈0.2, beta1符号反转)——提示理论假设与实证数据间存在重要偏差'
])

H(2,'7.2  理论贡献')
Bullets([
    '提出了Sigmoid(理论刚性)+GP(数据柔性)的混合建模框架——一种在认知建模领域具有方法论创新性的"理论引导、数据校准"方法',
    '通过Censor vs Drop系统对比为DDM应用中omission处理方法选择提供了基于真实数据的定量指导',
    '整合了"88人内部实验+44篇外部数据库+仿真验证"的多源证据链，提升了建模结论的可信度'
])

H(2,'7.3  方法论贡献')
Bullets([
    '开发了完整的DDM分析工具链(Docker HDDM+参数提取+Sigmoid校准+GP建模+LOCV+行为验证+候选推荐)',
    '首次将LAN+OPN框架应用于SMT范式，验证了Leng et al.(2025)方法论的跨范式可迁移性',
    'SPE Database标准化框架为后续跨研究数据整合提供了基础设施'
])

H(2,'7.4  局限与未来方向')
Bullets([
    '样本量限制：仅6-8个设计条件限制了GP泛化能力——未来需新增实验条件（如P=35,60在T=80/W=800上的验证）',
    'Sigmoid函数形式：base_scale_a触及搜索上界——a的参数化可能需要更灵活的函数形式',
    'DDM模型选择：当前使用恒定边界DDM——未来可考虑塌缩边界模型(ANGLE/WEIBULL)',
    'NonMatching试次：当前仅分析Matching试次——NonMatching的决策机制可能提供额外约束信息',
    '数据采集：88人数据由课题组前期采集，非本人独立收集——但方法论框架、分析工具链和外部验证为独立贡献'
])

H(2,'7.5  结论')
P('本研究构建了一个从实验设计参数(P,T,W)到心理决策过程(DDM参数v,a,t,z)再到行为输出(RT,ACC,Omission)的完整映射框架。通过Sigmoid理论先验与GP数据驱动的结合，模型能够在行为层面高度精确地重建真实被试的SPE模式。Omission敏感性分析和外部数据库验证为模型的可靠性和泛化性提供了多层次的实证支持。本研究为SPE实验设计的量化优化提供了方法论基础。')
doc.add_page_break()

# ═══ 参考文献 ═══
H(1,'参考文献')
refs = [
    'Dutilh, G., Krypotos, A.-M., & Wagenmakers, E.-J. (2011). Task-related versus stimulus-specific practice. Experimental Psychology, 58(6), 434-442.',
    'Leng, X., Fengler, A., Shenhav, A., & Frank, M. J. (2025). The Perils of Omitting Omissions when Modeling Evidence Accumulation. In Prep.',
    'Luce, R. D. (1986). Response Times: Their Role in Inferring Elementary Mental Organization. Oxford University Press.',
    'Matzke, D., & Wagenmakers, E.-J. (2009). Psychological interpretation of the ex-Gaussian and shifted Wald parameters. Psychonomic Bulletin & Review, 16(5), 798-817.',
    'Myung, J. I., Cavagnaro, D. R., & Pitt, M. A. (2013). A tutorial on adaptive design optimization. Journal of Mathematical Psychology, 57(3), 53-67.',
    'Ratcliff, R. (1978). A theory of memory retrieval. Psychological Review, 85(2), 59-108.',
    'Ratcliff, R., & McKoon, G. (2008). The diffusion decision model. Neural Computation, 20(4), 873-922.',
    'Schulz, E., Speekenbrink, M., & Krause, A. (2018). A tutorial on GP regression. Journal of Mathematical Psychology, 85, 1-16.',
    'Sui, J., He, X., & Humphreys, G. W. (2012). Perceptual effects of social salience. JEP:HPP, 38(5), 1105-1117.',
    'Tran, N. H., van Maanen, L., Heathcote, A., & Matzke, D. (2021). Systematic Parameter Reviews in Cognitive Modeling. Frontiers in Psychology, 11, 608287.',
    'Voss, A., Rothermund, K., & Voss, J. (2004). Interpreting the parameters of the diffusion model. Memory & Cognition, 32(7), 1206-1220.',
    'Cai, Z., Wang, Q., Liu, Z., Sui, J., & Hu, C.-P. (2026). The Self-Prioritization Effect Database with Standardized Meta-Data. In Prep.',
]
for i, ref in enumerate(refs):
    doc.add_paragraph(f'[{i+1}] {ref}', style='List Number')

# ═══ 附录 ═══
doc.add_page_break()
H(1,'附录')
H(2,'附录A：Sigmoid参数全表与可优化性分析')
P('（引用5_Reference/RoadMap.md Phase 0.4及理论逻辑链评估与修正建议.md）')
H(2,'附录B：DDM参数参考手册')
P('（引用5_Reference/RoadMap.md附录章节）')
H(2,'附录C：Omission建模可行性分析与LAN+OPN作战计划')
P('（引用5_Reference/Omission建模可行性分析.md, Omission_v2.md, Omission_LAN_OPN_实施作战计划.md）')
H(2,'附录D：SPE Database标准化框架')
P('（引用SPE_Database_README.md及SPE_数据库_v15.3.md）')
H(2,'附录E：项目代码与数据可复现性说明')
P('代码位于1_Code/，数据位于2_Data/，图表位于3_Figures/。计划在论文发表前上传至OSF/GitHub。')

doc.save(OUT)
print(f'✅ 毕业论文大纲v2已保存到: {OUT}')
